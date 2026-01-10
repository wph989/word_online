"""
DOCX 文档解析器

使用 python-docx 库解析 DOCX 文件结构，提取文档内容和格式信息
"""

import io
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
from docx import Document
from docx.shared import Pt, Twips, Emu, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell


# ============ 数据类定义 ============

@dataclass
class RunFormat:
    """Run 格式信息"""
    bold: bool = False
    italic: bool = False
    underline: bool = False
    strike: bool = False
    superscript: bool = False
    subscript: bool = False
    color: Optional[str] = None
    font_size: Optional[str] = None
    font_family: Optional[str] = None
    font_family_east_asia: Optional[str] = None
    background_color: Optional[str] = None
    small_caps: bool = False
    all_caps: bool = False


@dataclass
class RunData:
    """Run 数据"""
    text: str
    format: RunFormat
    start: int = 0
    end: int = 0


@dataclass
class ParagraphFormat:
    """段落格式信息"""
    alignment: Optional[str] = None
    first_line_indent: Optional[str] = None
    left_indent: Optional[str] = None
    right_indent: Optional[str] = None
    space_before: Optional[str] = None
    space_after: Optional[str] = None
    line_spacing: Optional[str] = None
    line_spacing_rule: Optional[str] = None


@dataclass
class ListInfo:
    """列表信息"""
    list_type: str  # "bullet" or "ordered"
    level: int
    num_id: Optional[int] = None
    start: Optional[int] = None


@dataclass  
class DocxParagraph:
    """解析后的段落数据"""
    text: str
    runs: List[RunData]
    format: ParagraphFormat
    heading_level: Optional[int] = None
    list_info: Optional[ListInfo] = None
    style_name: Optional[str] = None


@dataclass
class CellData:
    """单元格数据"""
    row: int
    col: int
    text: str
    runs: List[RunData]
    rowspan: int = 1
    colspan: int = 1
    width: Optional[str] = None
    background_color: Optional[str] = None
    vertical_align: Optional[str] = None
    text_align: Optional[str] = None  # 文本对齐方式
    border_top: Optional[Dict] = None
    border_bottom: Optional[Dict] = None
    border_left: Optional[Dict] = None
    border_right: Optional[Dict] = None


@dataclass
class DocxTable:
    """解析后的表格数据"""
    rows: int
    cols: int
    cells: List[CellData]
    column_widths: List[str] = field(default_factory=list)
    table_width: Optional[str] = None
    alignment: Optional[str] = None


@dataclass
class DocxImage:
    """解析后的图片数据"""
    rId: str
    width: Optional[str] = None
    height: Optional[str] = None
    alt: Optional[str] = None
    image_bytes: Optional[bytes] = None
    content_type: Optional[str] = None
    filename: Optional[str] = None


@dataclass
class PageSettings:
    """页面设置"""
    page_width: Optional[str] = None
    page_height: Optional[str] = None
    margin_top: Optional[int] = None
    margin_bottom: Optional[int] = None
    margin_left: Optional[int] = None
    margin_right: Optional[int] = None
    orientation: str = "portrait"


@dataclass
class DocxElement:
    """文档元素（段落、表格、图片的统一包装）"""
    element_type: str  # "paragraph", "table", "image"
    data: Any  # DocxParagraph, DocxTable, or DocxImage


@dataclass
class DocxParseResult:
    """解析结果"""
    elements: List[DocxElement]
    page_settings: PageSettings
    images: Dict[str, DocxImage]


# ============ 单位转换工具 ============

def twips_to_pt(twips: int) -> float:
    """Twips to points (1 pt = 20 twips)"""
    if twips is None:
        return None
    return twips / 20


def emu_to_pt(emu: int) -> float:
    """EMU to points (1 pt = 12700 EMU)"""
    if emu is None:
        return None
    return emu / 12700


def format_pt(value: float) -> str:
    """格式化为 pt 字符串"""
    if value is None:
        return None
    return f"{round(value, 1)}pt"


def format_px(pt_value: float, dpi: int = 96) -> str:
    """pt 转 px 字符串"""
    if pt_value is None:
        return None
    px = pt_value * dpi / 72
    return f"{round(px)}px"


# ============ 解析器实现 ============

class DocxParser:
    """
    DOCX 文档解析器
    
    解析 DOCX 文件结构，提取：
    1. 段落和标题
    2. 表格（含合并单元格）
    3. 图片
    4. 文本格式
    5. 段落格式
    6. 页面设置
    """
    
    def __init__(self, file_content: bytes):
        """
        初始化解析器
        
        Args:
            file_content: DOCX 文件的字节内容
        """
        self.doc = Document(io.BytesIO(file_content))
        self.images: Dict[str, DocxImage] = {}
        
    def parse(self) -> DocxParseResult:
        """
        解析文档
        
        Returns:
            DocxParseResult: 包含元素列表、页面设置和图片的解析结果
        """
        # 1. 提取图片
        self._extract_all_images()
        
        # 2. 解析文档元素
        elements = self._parse_body()
        
        # 3. 提取页面设置
        page_settings = self._get_page_settings()
        
        return DocxParseResult(
            elements=elements,
            page_settings=page_settings,
            images=self.images
        )
    
    def _parse_body(self) -> List[DocxElement]:
        """解析文档主体"""
        elements = []
        
        for element in self.doc.element.body:
            tag = element.tag.split('}')[-1]  # 去掉命名空间
            
            if tag == 'p':
                # 段落
                para = Paragraph(element, self.doc)
                parsed = self._parse_paragraph(para)
                if parsed:
                    elements.append(DocxElement(
                        element_type="paragraph",
                        data=parsed
                    ))
            elif tag == 'tbl':
                # 表格
                table = Table(element, self.doc)
                parsed = self._parse_table(table)
                if parsed:
                    elements.append(DocxElement(
                        element_type="table",
                        data=parsed
                    ))
        
        return elements
    
    def _parse_paragraph(self, para: Paragraph) -> Optional[DocxParagraph]:
        """解析段落"""
        # 检查是否包含图片
        inline_shapes = para._element.findall('.//a:blip', 
            {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
        
        # 如果段落只有图片没有文字，返回 None（图片单独处理）
        text = para.text
        
        # 提取所有 Runs
        runs = self._parse_runs(para)
        
        # 提取段落格式
        para_format = self._parse_paragraph_format(para)
        
        # 检测标题级别
        heading_level = self._get_heading_level(para)
        
        # 检测列表信息
        list_info = self._get_list_info(para)
        
        # 检查段落中的图片
        drawing_elements = para._element.findall('.//w:drawing', 
            {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        
        return DocxParagraph(
            text=text,
            runs=runs,
            format=para_format,
            heading_level=heading_level,
            list_info=list_info,
            style_name=para.style.name if para.style else None
        )
    
    def _parse_runs(self, para: Paragraph) -> List[RunData]:
        """解析段落中的所有 Run"""
        runs = []
        current_pos = 0
        
        for run in para.runs:
            text = run.text
            if not text:
                continue
            
            start = current_pos
            end = current_pos + len(text)
            
            run_format = self._parse_run_format(run)
            
            runs.append(RunData(
                text=text,
                format=run_format,
                start=start,
                end=end
            ))
            
            current_pos = end
        
        return runs
    
    def _parse_run_format(self, run) -> RunFormat:
        """解析 Run 格式"""
        font = run.font
        
        # 颜色处理
        color = None
        if font.color and font.color.rgb:
            rgb = font.color.rgb
            color = f"rgb({rgb[0]}, {rgb[1]}, {rgb[2]})" if isinstance(rgb, tuple) else f"#{rgb}"
        
        # 字号处理
        font_size = None
        if font.size:
            font_size = format_pt(emu_to_pt(font.size))
        
        # 高亮/背景色处理
        background_color = None
        if font.highlight_color:
            # highlight_color 是枚举值
            background_color = str(font.highlight_color).lower()
        
        return RunFormat(
            bold=bool(run.bold),
            italic=bool(run.italic),
            underline=bool(run.underline),
            strike=bool(font.strike) if font.strike is not None else False,
            superscript=bool(font.superscript),
            subscript=bool(font.subscript),
            color=color,
            font_size=font_size,
            font_family=font.name,
            font_family_east_asia=font.name_eastasia if hasattr(font, 'name_eastasia') else None,
            background_color=background_color,
            small_caps=bool(font.small_caps) if font.small_caps is not None else False,
            all_caps=bool(font.all_caps) if font.all_caps is not None else False
        )
    
    def _parse_paragraph_format(self, para: Paragraph) -> ParagraphFormat:
        """解析段落格式"""
        pf = para.paragraph_format
        
        # 对齐方式
        alignment = None
        if para.alignment is not None:
            alignment_map = {
                WD_ALIGN_PARAGRAPH.LEFT: "left",
                WD_ALIGN_PARAGRAPH.CENTER: "center",
                WD_ALIGN_PARAGRAPH.RIGHT: "right",
                WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
                WD_ALIGN_PARAGRAPH.DISTRIBUTE: "justify",  # 分散对齐映射为两端对齐
            }
            alignment = alignment_map.get(para.alignment, "left")
        
        # 首行缩进
        first_line_indent = None
        if pf.first_line_indent:
            pt_val = emu_to_pt(pf.first_line_indent)
            first_line_indent = format_pt(pt_val)
        
        # 左右缩进
        left_indent = format_pt(emu_to_pt(pf.left_indent)) if pf.left_indent else None
        right_indent = format_pt(emu_to_pt(pf.right_indent)) if pf.right_indent else None
        
        # 段前段后间距
        space_before = format_pt(emu_to_pt(pf.space_before)) if pf.space_before else None
        space_after = format_pt(emu_to_pt(pf.space_after)) if pf.space_after else None
        
        # 行间距
        line_spacing = None
        line_spacing_rule = None
        if pf.line_spacing is not None:
            if pf.line_spacing_rule == WD_LINE_SPACING.SINGLE:
                line_spacing = "1"
            elif pf.line_spacing_rule == WD_LINE_SPACING.ONE_POINT_FIVE:
                line_spacing = "1.5"
            elif pf.line_spacing_rule == WD_LINE_SPACING.DOUBLE:
                line_spacing = "2"
            elif pf.line_spacing_rule == WD_LINE_SPACING.MULTIPLE:
                line_spacing = str(pf.line_spacing)
            elif pf.line_spacing_rule == WD_LINE_SPACING.EXACTLY:
                line_spacing = format_pt(emu_to_pt(pf.line_spacing))
                line_spacing_rule = "exact"
            elif pf.line_spacing_rule == WD_LINE_SPACING.AT_LEAST:
                line_spacing = format_pt(emu_to_pt(pf.line_spacing))
                line_spacing_rule = "atLeast"
        
        return ParagraphFormat(
            alignment=alignment,
            first_line_indent=first_line_indent,
            left_indent=left_indent,
            right_indent=right_indent,
            space_before=space_before,
            space_after=space_after,
            line_spacing=line_spacing,
            line_spacing_rule=line_spacing_rule
        )
    
    def _get_heading_level(self, para: Paragraph) -> Optional[int]:
        """
        识别标题级别
        
        支持：
        1. 内置标题样式 (Heading 1, Heading 2, ...)
        2. 大纲级别
        """
        if para.style is None:
            return None
        
        style_name = para.style.name
        
        # 方式1: 内置样式名
        if style_name.startswith('Heading'):
            try:
                level = int(style_name.split()[-1])
                if 1 <= level <= 6:
                    return level
            except (ValueError, IndexError):
                pass
        
        # 中文样式名
        heading_cn_map = {
            '标题 1': 1, '标题 2': 2, '标题 3': 3,
            '标题 4': 4, '标题 5': 5, '标题 6': 6,
            '标题1': 1, '标题2': 2, '标题3': 3,
            '标题4': 4, '标题5': 5, '标题6': 6,
        }
        if style_name in heading_cn_map:
            return heading_cn_map[style_name]
        
        # 方式2: 大纲级别 (通过 XML 检查)
        try:
            pPr = para._element.pPr
            if pPr is not None:
                outlineLvl = pPr.find(qn('w:outlineLvl'))
                if outlineLvl is not None:
                    val = outlineLvl.get(qn('w:val'))
                    if val is not None:
                        level = int(val) + 1  # 大纲级别从 0 开始
                        if 1 <= level <= 6:
                            return level
        except Exception:
            pass
        
        return None
    
    def _get_list_info(self, para: Paragraph) -> Optional[ListInfo]:
        """获取列表信息"""
        try:
            pPr = para._element.pPr
            if pPr is None:
                return None
            
            numPr = pPr.find(qn('w:numPr'))
            if numPr is None:
                return None
            
            # 获取 numId 和 ilvl
            numId_elem = numPr.find(qn('w:numId'))
            ilvl_elem = numPr.find(qn('w:ilvl'))
            
            if numId_elem is None:
                return None
            
            num_id = int(numId_elem.get(qn('w:val')))
            level = int(ilvl_elem.get(qn('w:val'))) if ilvl_elem is not None else 0
            
            # 判断列表类型（需要从 numbering.xml 获取）
            list_type = self._get_list_type(num_id, level)
            
            return ListInfo(
                list_type=list_type,
                level=level,
                num_id=num_id
            )
        except Exception:
            return None
    
    def _get_list_type(self, num_id: int, level: int) -> str:
        """获取列表类型（有序/无序）"""
        try:
            numbering_part = self.doc.part.numbering_part
            if numbering_part is None:
                return "bullet"
            
            numbering_xml = numbering_part._element
            
            # 查找 num 元素
            for num in numbering_xml.findall(qn('w:num')):
                if num.get(qn('w:numId')) == str(num_id):
                    abstractNumId = num.find(qn('w:abstractNumId'))
                    if abstractNumId is not None:
                        abstract_id = abstractNumId.get(qn('w:val'))
                        
                        # 查找 abstractNum
                        for abstractNum in numbering_xml.findall(qn('w:abstractNum')):
                            if abstractNum.get(qn('w:abstractNumId')) == abstract_id:
                                # 查找对应级别的 lvl
                                for lvl in abstractNum.findall(qn('w:lvl')):
                                    if lvl.get(qn('w:ilvl')) == str(level):
                                        numFmt = lvl.find(qn('w:numFmt'))
                                        if numFmt is not None:
                                            fmt = numFmt.get(qn('w:val'))
                                            if fmt == 'bullet':
                                                return "bullet"
                                            else:
                                                return "ordered"
            return "bullet"
        except Exception:
            return "bullet"
    
    def _parse_table(self, table: Table) -> DocxTable:
        """解析表格"""
        rows = len(table.rows)
        cols = len(table.columns)
        
        cells = []
        column_widths = self._extract_column_widths(table)
        
        # 用于跟踪合并单元格
        merge_map = {}  # (row, col) -> (master_row, master_col)
        
        for row_idx, row in enumerate(table.rows):
            col_idx = 0
            for cell in row.cells:
                # 跳过已被合并的单元格
                while (row_idx, col_idx) in merge_map:
                    col_idx += 1
                
                if col_idx >= cols:
                    break
                
                # 获取合并信息
                tc = cell._tc
                rowspan, colspan = self._get_cell_span(tc)
                
                # 标记被合并的单元格
                for dr in range(rowspan):
                    for dc in range(colspan):
                        if dr > 0 or dc > 0:
                            merge_map[(row_idx + dr, col_idx + dc)] = (row_idx, col_idx)
                
                # 获取单元格样式
                cell_style = self._parse_cell_style(cell)
                
                # 解析单元格内容
                cell_text = cell.text
                cell_runs = []
                cell_text_align = None
                
                # 从单元格段落获取文本对齐方式
                # 优先使用第一个非空段落的对齐方式
                if cell.paragraphs:
                    for para in cell.paragraphs:
                        if para.text.strip() and para.alignment is not None:
                            alignment_map = {
                                WD_ALIGN_PARAGRAPH.LEFT: "left",
                                WD_ALIGN_PARAGRAPH.CENTER: "center",
                                WD_ALIGN_PARAGRAPH.RIGHT: "right",
                                WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
                                WD_ALIGN_PARAGRAPH.DISTRIBUTE: "justify",
                            }
                            cell_text_align = alignment_map.get(para.alignment)
                            break
                    
                    # 如果没有找到，使用第一个段落的对齐方式（即使为空）
                    if cell_text_align is None and cell.paragraphs[0].alignment is not None:
                        alignment_map = {
                            WD_ALIGN_PARAGRAPH.LEFT: "left",
                            WD_ALIGN_PARAGRAPH.CENTER: "center",
                            WD_ALIGN_PARAGRAPH.RIGHT: "right",
                            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
                            WD_ALIGN_PARAGRAPH.DISTRIBUTE: "justify",
                        }
                        cell_text_align = alignment_map.get(cell.paragraphs[0].alignment)
                
                for para in cell.paragraphs:
                    cell_runs.extend(self._parse_runs(para))
                
                cells.append(CellData(
                    row=row_idx,
                    col=col_idx,
                    text=cell_text,
                    runs=cell_runs,
                    rowspan=rowspan,
                    colspan=colspan,
                    width=cell_style.get('width'),
                    background_color=cell_style.get('backgroundColor'),
                    vertical_align=cell_style.get('verticalAlign'),
                    text_align=cell_text_align,
                    border_top=cell_style.get('borderTop'),
                    border_bottom=cell_style.get('borderBottom'),
                    border_left=cell_style.get('borderLeft'),
                    border_right=cell_style.get('borderRight')
                ))
                
                col_idx += colspan
        
        # 表格对齐方式
        alignment = None
        if table.alignment:
            alignment_map = {
                WD_TABLE_ALIGNMENT.LEFT: "left",
                WD_TABLE_ALIGNMENT.CENTER: "center",
                WD_TABLE_ALIGNMENT.RIGHT: "right",
            }
            alignment = alignment_map.get(table.alignment)
        
        return DocxTable(
            rows=rows,
            cols=cols,
            cells=cells,
            column_widths=column_widths,
            alignment=alignment
        )
    
    def _extract_column_widths(self, table: Table) -> List[str]:
        """提取表格列宽"""
        widths = []
        
        try:
            # 方法1: 从 tblGrid 获取
            tblGrid = table._tbl.find(qn('w:tblGrid'))
            if tblGrid is not None:
                for gridCol in tblGrid.findall(qn('w:gridCol')):
                    w = gridCol.get(qn('w:w'))
                    if w:
                        # 转换 twips 到 pt
                        pt_val = twips_to_pt(int(w))
                        widths.append(format_pt(pt_val))
            
            # 方法2: 如果没有 tblGrid，从第一行单元格获取
            if not widths and len(table.rows) > 0:
                for cell in table.rows[0].cells:
                    if cell.width:
                        pt_val = emu_to_pt(cell.width)
                        widths.append(format_pt(pt_val))
        except Exception:
            pass
        
        return widths
    
    def _get_cell_span(self, tc) -> Tuple[int, int]:
        """获取单元格的 rowspan 和 colspan"""
        rowspan = 1
        colspan = 1
        
        tcPr = tc.tcPr
        if tcPr is not None:
            # colspan (gridSpan)
            gridSpan = tcPr.find(qn('w:gridSpan'))
            if gridSpan is not None:
                val = gridSpan.get(qn('w:val'))
                if val:
                    colspan = int(val)
            
            # rowspan (vMerge)
            vMerge = tcPr.find(qn('w:vMerge'))
            if vMerge is not None:
                val = vMerge.get(qn('w:val'))
                if val == 'restart':
                    # 这是合并的起始单元格，需要计算 rowspan
                    rowspan = self._calculate_rowspan(tc)
        
        return rowspan, colspan
    
    def _calculate_rowspan(self, start_tc) -> int:
        """计算单元格的 rowspan"""
        # 简化处理：返回 1，实际合并信息已在 _get_cell_span 中处理
        # 完整实现需要遍历后续行检查 vMerge
        return 1
    
    def _parse_cell_style(self, cell: _Cell) -> Dict[str, Any]:
        """解析单元格样式"""
        style = {}
        
        tc = cell._tc
        tcPr = tc.tcPr
        
        if tcPr is None:
            return style
        
        # 背景色
        shd = tcPr.find(qn('w:shd'))
        if shd is not None:
            fill = shd.get(qn('w:fill'))
            if fill and fill.lower() != 'auto':
                style['backgroundColor'] = f"#{fill}"
        
        # 垂直对齐
        vAlign = tcPr.find(qn('w:vAlign'))
        if vAlign is not None:
            val = vAlign.get(qn('w:val'))
            align_map = {'top': 'top', 'center': 'middle', 'bottom': 'bottom'}
            style['verticalAlign'] = align_map.get(val, 'top')
        
        # 宽度
        tcW = tcPr.find(qn('w:tcW'))
        if tcW is not None:
            w = tcW.get(qn('w:w'))
            w_type = tcW.get(qn('w:type'))
            if w:
                if w_type == 'pct':
                    style['width'] = f"{int(w) / 50}%"  # 转换为百分比
                else:
                    style['width'] = format_pt(twips_to_pt(int(w)))
        
        # 边框 (简化处理)
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is not None:
            for side in ['top', 'bottom', 'left', 'right']:
                border = tcBorders.find(qn(f'w:{side}'))
                if border is not None:
                    style[f'border{side.capitalize()}'] = self._parse_border(border)
        
        return style
    
    def _parse_border(self, border_elem) -> Dict[str, str]:
        """解析边框样式"""
        return {
            'width': format_pt(int(border_elem.get(qn('w:sz'), 0)) / 8),  # 1/8 pt
            'style': border_elem.get(qn('w:val'), 'single'),
            'color': f"#{border_elem.get(qn('w:color'), '000000')}"
        }
    
    def _extract_all_images(self):
        """提取文档中的所有图片"""
        try:
            for rel in self.doc.part.rels.values():
                if "image" in rel.target_ref.lower():
                    try:
                        image_part = rel.target_part
                        self.images[rel.rId] = DocxImage(
                            rId=rel.rId,
                            image_bytes=image_part.blob,
                            content_type=image_part.content_type,
                            filename=rel.target_ref.split('/')[-1]
                        )
                    except Exception:
                        pass
        except Exception:
            pass
    
    def _get_page_settings(self) -> PageSettings:
        """获取页面设置"""
        settings = PageSettings()
        
        try:
            section = self.doc.sections[0] if self.doc.sections else None
            if section:
                settings.page_width = format_pt(emu_to_pt(section.page_width)) if section.page_width else None
                settings.page_height = format_pt(emu_to_pt(section.page_height)) if section.page_height else None
                settings.margin_top = round(emu_to_pt(section.top_margin)) if section.top_margin else None
                settings.margin_bottom = round(emu_to_pt(section.bottom_margin)) if section.bottom_margin else None
                settings.margin_left = round(emu_to_pt(section.left_margin)) if section.left_margin else None
                settings.margin_right = round(emu_to_pt(section.right_margin)) if section.right_margin else None
                settings.orientation = "landscape" if section.page_width > section.page_height else "portrait"
        except Exception:
            pass
        
        return settings
