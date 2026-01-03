"""
Word 文档导出服务
将 Content 和 StyleSheet 导出为 .docx 文件

功能特性:
1. 支持段落、标题、表格、图片、代码块、分割线
2. 支持丰富的文本标记(加粗、斜体、颜色、字号、字体等)
3. 支持表格合并单元格和单元格样式
4. 支持列表(有序、无序)
5. 支持行高、缩进、对齐等段落样式
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, Any, List, Optional, Tuple, Union
import io
import re
import base64
import requests
from io import BytesIO
import logging

from app.config.editor_defaults import (
    get_default_font_family,
    get_default_font_size_pt,
    get_default_line_height
)

logger = logging.getLogger(__name__)


class DocxExporter:
    """
    Word 文档导出器
    将结构化的 Content 和 StyleSheet 转换为 Word 文档
    """
    
    def __init__(
        self, 
        content: Dict[str, Any], 
        stylesheet: Dict[str, Any],
        document_settings: Optional[Dict[str, Any]] = None
    ):
        """
        初始化导出器
        
        Args:
            content: Content JSON
            stylesheet: StyleSheet JSON
            document_settings: 文档配置 (页边距、标题样式等)
        """
        self.content = content
        self.stylesheet = stylesheet
        self.document_settings = document_settings or {}
        self.doc = Document()
        
        # 应用页边距配置
        self._apply_page_margins()
        
        # 配置标题样式模板
        self._configure_heading_styles()
        
        self.style_map = self._build_style_map()
        self.cell_style_map = self._build_cell_style_map()
        
        logger.info(f"初始化 DocxExporter, Blocks 数量: {len(content.get('blocks', []))}")
        if document_settings:
            logger.info(f"应用文档配置: 页边距={document_settings.get('margin_top', 2.54)}cm(上)")
    
    def _apply_page_margins(self):
        """应用页边距配置到文档"""
        if not self.document_settings:
            return
        
        # 获取页边距配置 (单位: cm)
        margin_top = self.document_settings.get('margin_top', 2.54)
        margin_bottom = self.document_settings.get('margin_bottom', 2.54)
        margin_left = self.document_settings.get('margin_left', 3.17)
        margin_right = self.document_settings.get('margin_right', 3.17)
        
        # 应用到所有节（Section）
        for section in self.doc.sections:
            section.top_margin = Cm(margin_top)
            section.bottom_margin = Cm(margin_bottom)
            section.left_margin = Cm(margin_left)
            section.right_margin = Cm(margin_right)
        
        logger.info(f"✅ 应用页边距: 上={margin_top}cm, 下={margin_bottom}cm, 左={margin_left}cm, 右={margin_right}cm")
    
    def _configure_heading_styles(self):
        """配置 Word 的标题样式模板"""
        if not self.document_settings or 'heading_styles' not in self.document_settings:
            return
        
        heading_styles_config = self.document_settings['heading_styles']
        
        for level in range(1, 7):  # Word 支持 Heading 1-6
            h_key = f"h{level}"
            if h_key not in heading_styles_config:
                continue
            
            config = heading_styles_config[h_key]
            style_name = f'Heading {level}'
            
            try:
                # 获取或创建标题样式
                if style_name in self.doc.styles:
                    style = self.doc.styles[style_name]
                else:
                    logger.warning(f"样式 {style_name} 不存在,跳过配置")
                    continue
                
                # 配置字体
                if config.get('fontFamily'):
                    style.font.name = config['fontFamily']
                    # 设置中文字体
                    style.element.rPr.rFonts.set(qn('w:eastAsia'), config['fontFamily'])
                
                # 配置字号 (pt)
                if config.get('fontSize'):
                    style.font.size = Pt(config['fontSize'])
                
                # 配置颜色
                if config.get('color'):
                    rgb = self._parse_color(config['color'])
                    if rgb:
                        style.font.color.rgb = RGBColor(*rgb)
                
                # 配置加粗
                if config.get('fontWeight') == 'bold':
                    style.font.bold = True
                
                # 配置段落间距
                if config.get('marginTop'):
                    style.paragraph_format.space_before = Pt(config['marginTop'])
                if config.get('marginBottom'):
                    style.paragraph_format.space_after = Pt(config['marginBottom'])
                
                logger.info(f"✅ 配置标题样式: {style_name} - 字体={config.get('fontFamily')}, 字号={config.get('fontSize')}pt")
                
            except Exception as e:
                logger.error(f"配置标题样式 {style_name} 失败: {e}", exc_info=True)
    
    def export(self) -> io.BytesIO:
        """
        执行导出,生成 Word 文档
        
        Returns:
            包含 .docx 文件的 BytesIO 对象
        """
        try:
            # 处理所有 Block
            for idx, block in enumerate(self.content.get("blocks", [])):
                try:
                    self._process_block(block)
                except Exception as e:
                    logger.error(f"处理 Block {idx} 失败: {e}", exc_info=True)
                    # 添加错误占位符
                    para = self.doc.add_paragraph(f"[导出错误: {block.get('type', 'unknown')}]")
                    para.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            
            # 保存到内存
            file_stream = io.BytesIO()
            self.doc.save(file_stream)
            file_stream.seek(0)
            
            logger.info(f"导出成功,文件大小: {file_stream.getbuffer().nbytes} 字节")
            return file_stream
            
        except Exception as e:
            logger.error(f"导出失败: {e}", exc_info=True)
            raise
    
    def _build_style_map(self) -> Dict[str, Dict[str, Any]]:
        """构建 Block ID 到样式的映射"""
        style_map = {}
        
        for rule in self.stylesheet.get("rules", []):
            target = rule.get("target", {})
            block_ids = target.get("blockIds", [])
            styles = rule.get("style", {})
            
            # 只处理 Block 级别的样式
            if target.get("blockType") in ["paragraph", "heading", "table"]:
                for block_id in block_ids:
                    if block_id not in style_map:
                        style_map[block_id] = {}
                    style_map[block_id].update(styles)
        
        return style_map
    
    def _build_cell_style_map(self) -> Dict[str, Dict[str, Any]]:
        """构建单元格 ID 到样式的映射"""
        cell_style_map = {}
        
        for rule in self.stylesheet.get("rules", []):
            target = rule.get("target", {})
            
            # 处理表格单元格样式 - 使用 blockIds
            if target.get("blockType") == "tableCell":
                block_ids = target.get("blockIds", [])
                styles = rule.get("style", {})
                
                for block_id in block_ids:
                    if block_id not in cell_style_map:
                        cell_style_map[block_id] = {}
                    cell_style_map[block_id].update(styles)
        
        return cell_style_map
    
    def _process_block(self, block: Dict[str, Any]):
        """
        处理单个 Block
        
        Args:
            block: Block 字典
        """
        block_type = block.get("type")
        
        if block_type == "paragraph":
            self._add_paragraph(block)
        elif block_type == "heading":
            self._add_heading(block)
        elif block_type == "table":
            self._add_table(block)
        elif block_type == "image":
            self._add_image(block)
        elif block_type == "code":
            self._add_code(block)
        elif block_type == "divider":
            self._add_divider(block)
        else:
            logger.warning(f"未知的 Block 类型: {block_type}")
    
    def _add_paragraph(self, block: Dict[str, Any]):
        """添加段落"""
        text = block.get("text", "")
        marks = block.get("marks", [])
        block_id = block.get("id")
        attrs = block.get("attrs") or {}  # 处理 null 值
        
        # 创建段落
        para = self.doc.add_paragraph()
        
        # 应用列表样式
        list_type = attrs.get("listType")
        list_level = attrs.get("listLevel", 0)
        
        if list_type == "bullet":
            para.style = 'List Bullet'
            # 设置列表层级
            if list_level > 0:
                para.paragraph_format.left_indent = Inches(0.5 * (list_level + 1))
        elif list_type == "ordered":
            para.style = 'List Number'
            if list_level > 0:
                para.paragraph_format.left_indent = Inches(0.5 * (list_level + 1))
        
        # 添加文本和标记
        self._add_text_with_marks(para, text, marks)
        
        # 应用段落样式
        self._apply_paragraph_style(para, block_id)
    
    def _add_heading(self, block: Dict[str, Any]):
        """添加标题
        
        标题格式处理规则:
        1. 标题样式模板在 _configure_heading_styles() 中已配置到 Word 文档
        2. 创建的标题段落会自动继承对应的样式模板(Heading 1-6)
        3. 如果文本有 mark 格式,在样式模板基础上叠加这些格式
        4. mark 格式只影响当前标题文本,不修改样式模板本身
        """
        level = block.get("level", 1)
        text = block.get("text", "")
        marks = block.get("marks", [])
        block_id = block.get("id")
        
        # 创建标题段落 (会自动应用对应的 Heading 样式模板)
        para = self.doc.add_heading(level=level)
        para.text = ""  # 清空默认文本
        # 显式清除所有 runs,防止保留一个空 run
        for run in list(para.runs):
            run._element.getparent().remove(run._element)
        
        # 添加文本和标记
        # inherit_paragraph_style=True 表示 run 继承段落样式模板,不应用默认字体和字号
        # marks 会在样式模板的基础上叠加(如局部加粗、变色等)
        self._add_text_with_marks(para, text, marks, inherit_paragraph_style=True)
        
        # 应用 StyleSheet 中的段落级样式(如对齐、行高等)
        self._apply_paragraph_style(para, block_id)
    
    def _add_table(self, block: Dict[str, Any]):
        """添加表格(支持合并单元格和单元格样式)"""
        table_data = block.get("data", {})
        rows = table_data.get("rows", 0)
        cols = table_data.get("cols", 0)
        cells = table_data.get("cells", [])
        merge_regions = table_data.get("mergeRegions", [])
        
        if rows == 0 or cols == 0:
            logger.warning("表格行数或列数为 0,跳过")
            return
        
        # 创建表格
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        # 构建单元格映射
        cell_map = {tuple(cell["cell"]): cell for cell in cells}
        
        # 填充单元格内容和样式
        for row_idx in range(rows):
            for col_idx in range(cols):
                cell_data = cell_map.get((row_idx, col_idx))
                if cell_data:
                    cell = table.cell(row_idx, col_idx)
                    content = cell_data.get("content", {})
                    text = content.get("text", "")
                    marks = content.get("marks", [])
                    
                    # 清空默认段落
                    cell.text = ""
                    para = cell.paragraphs[0]
                    
                    # 添加文本和标记
                    self._add_text_with_marks(para, text, marks)
                    
                    # 应用单元格样式 - 使用 cell-X-Y 格式的 ID
                    cell_id = f"cell-{row_idx}-{col_idx}"
                    self._apply_cell_style(cell, para, cell_id)
                    
                    # 为表头行（第一行）应用默认样式
                    if row_idx == 0:
                        self._apply_header_style(cell, para)
        
        # 处理合并单元格
        for region in merge_regions:
            try:
                start_row, start_col = region["start"]
                end_row, end_col = region["end"]
                
                # 合并单元格
                start_cell = table.cell(start_row, start_col)
                end_cell = table.cell(end_row, end_col)
                start_cell.merge(end_cell)
            except Exception as e:
                logger.error(f"合并单元格失败: {e}")
    
    def _add_image(self, block: Dict[str, Any]):
        """添加图片(支持本地路径、URL 和 base64)"""
        src = block.get("src", "")
        meta = block.get("meta", {})
        
        try:
            image_stream = None
            
            # 处理 base64 图片
            if src.startswith("data:image"):
                # 提取 base64 数据
                match = re.match(r'data:image/[^;]+;base64,(.+)', src)
                if match:
                    base64_data = match.group(1)
                    image_data = base64.b64decode(base64_data)
                    image_stream = BytesIO(image_data)
            
            # 处理 URL 图片
            elif src.startswith("http://") or src.startswith("https://"):
                response = requests.get(src, timeout=10)
                response.raise_for_status()
                image_stream = BytesIO(response.content)
            
            # 处理本地路径
            else:
                image_stream = src
            
            # 添加图片
            if image_stream:
                width = meta.get("width")
                if width and isinstance(width, (int, float)):
                    width_inches = Inches(width / 96)  # 假设 96 DPI
                    self.doc.add_picture(image_stream, width=width_inches)
                else:
                    self.doc.add_picture(image_stream)
                
                logger.info(f"成功添加图片: {src[:50]}...")
            else:
                raise ValueError("无法识别的图片源")
                
        except Exception as e:
            logger.error(f"添加图片失败: {e}")
            # 添加占位文本
            para = self.doc.add_paragraph(f"[图片加载失败: {src[:50]}...]")
            para.runs[0].italic = True
            para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    def _add_code(self, block: Dict[str, Any]):
        """添加代码块"""
        text = block.get("text", "")
        language = block.get("language", "")
        
        # 添加语言标签(如果有)
        if language:
            lang_para = self.doc.add_paragraph(f"[{language}]")
            lang_para.runs[0].font.size = Pt(9)
            lang_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        # 添加代码内容
        para = self.doc.add_paragraph(text)
        para.style = 'No Spacing'
        
        # 设置等宽字体和样式
        for run in para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        
        # 添加背景色(通过 shading)
        self._set_paragraph_shading(para, RGBColor(245, 245, 245))
    
    def _add_divider(self, block: Dict[str, Any]):
        """添加分割线"""
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        
        # 添加底部边框作为分割线
        pPr = para._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')  # 线宽
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')  # 黑色
        
        pBdr.append(bottom)
        pPr.append(pBdr)
    
    def _add_text_with_marks(self, para, text: str, marks: List[Dict[str, Any]], default_style: Optional[Dict[str, Any]] = None, filter_mark_types: Optional[List[str]] = None, inherit_paragraph_style: bool = False):
        """
        添加带标记的文本到段落
        
        Args:
            para: Word 段落对象
            text: 文本内容
            marks: 标记列表
            default_style: 默认样式 (如标题的字号、颜色等)
            filter_mark_types: 需要过滤的 mark 类型列表（这些 mark 将被忽略）
            inherit_paragraph_style: 是否继承段落样式(True 时不应用默认字体和字号)
        """
        if not text:
            return
        
        # 过滤 marks
        if filter_mark_types:
            marks = [m for m in marks if m.get("type") not in filter_mark_types]
        
        
        if not marks:
            run = para.add_run(text)
            self._apply_default_style_to_run(run, default_style, inherit_paragraph_style)
            return
        
        # 构建字符级别的标记映射
        char_marks = [[] for _ in range(len(text))]
        for mark in marks:
            start, end = mark["range"]
            for i in range(start, min(end, len(text))):
                char_marks[i].append(mark)
        
        # 合并相同标记的连续字符
        segments = []
        current_start = 0
        current_marks = char_marks[0] if char_marks else []
        
        for i in range(1, len(text)):
            if char_marks[i] != current_marks:
                segments.append((current_start, i, current_marks))
                current_start = i
                current_marks = char_marks[i]
        
        # 添加最后一段
        segments.append((current_start, len(text), current_marks))
        
        # 为每个段落创建 run 并应用标记
        for start, end, segment_marks in segments:
            run = para.add_run(text[start:end])
            
            # 1. 先应用默认样式 (如果有传入 default_style，则使用它；否则使用全局默认)
            self._apply_default_style_to_run(run, default_style, inherit_paragraph_style)
            
            # 2. 再应用字符级标记（会覆盖默认值）
            self._apply_marks_to_run(run, segment_marks)
    
    def _apply_default_style_to_run(self, run, default_style: Optional[Dict[str, Any]], inherit_paragraph_style: bool = False):
        """应用默认样式到 run
        
        Args:
            run: Word run 对象
            default_style: 默认样式字典
            inherit_paragraph_style: 是否继承段落样式(True 时不设置字体和字号)
        """
        # 如果继承段落样式(如标题),则不设置默认字体和字号
        # 让 run 继承段落样式模板的设置
        if not inherit_paragraph_style:
            # 基础默认值(仅用于普通段落)
            run.font.name = get_default_font_family()
            run.font.size = Pt(get_default_font_size_pt())
        
        if not default_style:
            return
            
        # 应用传入的覆盖样式
        if default_style.get("fontSize"):
            pt_size = default_style["fontSize"]
            run.font.size = Pt(pt_size)
        
        if default_style.get("color"):
            rgb = self._parse_color(default_style["color"])
            if rgb:
                run.font.color.rgb = RGBColor(*rgb)
        
        if default_style.get("fontFamily"):
            font_family = default_style["fontFamily"]
            run.font.name = font_family
            # 设置中文字体
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_family)

        if default_style.get("fontWeight") == "bold":
            run.bold = True
    
    
    def _apply_marks_to_run(self, run, marks: List[Dict[str, Any]]):
        """应用标记到 run"""
        for mark in marks:
            mark_type = mark["type"]
            
            if mark_type == "bold":
                run.bold = True
            elif mark_type == "italic":
                run.italic = True
            elif mark_type == "underline":
                run.underline = True
            elif mark_type == "strike":
                run.font.strike = True
            elif mark_type == "superscript":
                run.font.superscript = True
            elif mark_type == "subscript":
                run.font.subscript = True
            elif mark_type == "color":
                color_str = mark.get("value", "")
                rgb = self._parse_color(color_str)
                if rgb:
                    run.font.color.rgb = RGBColor(*rgb)
            elif mark_type == "backgroundColor":
                color_str = mark.get("value", "")
                rgb = self._parse_color(color_str)
                if rgb:
                    self._set_run_background(run, RGBColor(*rgb))
            elif mark_type == "fontSize":
                size_str = mark.get("value", "")
                size = self._parse_font_size(size_str)
                if size:
                    run.font.size = Pt(size)
            elif mark_type == "fontFamily":
                family = mark.get("value", "")
                if family:
                    # 清理字体名称：移除引号和多余空格
                    family = family.strip().strip('"').strip("'")
                    # 如果有多个字体（用逗号分隔），只取第一个
                    if ',' in family:
                        family = family.split(',')[0].strip().strip('"').strip("'")
                    
                    logger.debug(f"应用字体: {family}")
                    run.font.name = family
                    # 同时设置东亚字体（对中文很重要）
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), family)
    
    def _apply_paragraph_style(self, para, block_id: str):
        """
        应用段落样式
        
        Args:
            para: Word 段落对象
            block_id: Block ID
        """
        styles = self.style_map.get(block_id, {})
        
        # 文本对齐
        align = styles.get("textAlign")
        if align == "left":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align == "center":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == "justify":
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # 行高
        line_height = styles.get("lineHeight")
        if line_height:
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = line_height
        else:
            # 如果段落样式中没有设置行高，应用默认行高
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = get_default_line_height()
        
        # 缩进
        text_indent = styles.get("textIndent")
        if text_indent:
            para.paragraph_format.first_line_indent = Pt(text_indent)
        
        # 段落间距
        margin_top = styles.get("marginTop")
        if margin_top:
            para.paragraph_format.space_before = Pt(margin_top)
        
        margin_bottom = styles.get("marginBottom")
        if margin_bottom:
            para.paragraph_format.space_after = Pt(margin_bottom)
        
        # 注意：字号和字体不在段落级别处理，只在字符级别（marks）处理
        # 这样可以避免段落样式覆盖字符级样式
        
        # 颜色
        color = styles.get("color")
        if color:
            rgb = self._parse_color(color)
            if rgb:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(*rgb)
        
        # 背景色
        background_color = styles.get("backgroundColor")
        if background_color:
            rgb = self._parse_color(background_color)
            if rgb:
                self._set_paragraph_shading(para, RGBColor(*rgb))
    
    def _apply_cell_style(self, cell, para, cell_id: str):
        """应用单元格样式"""
        styles = self.cell_style_map.get(cell_id, {})
        
        # 背景色
        background_color = styles.get("backgroundColor")
        if background_color:
            rgb = self._parse_color(background_color)
            if rgb:
                self._set_cell_background(cell, RGBColor(*rgb))
        
        # 文本对齐
        text_align = styles.get("textAlign")
        if text_align == "left":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif text_align == "center":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif text_align == "right":
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # 垂直对齐
        vertical_align = styles.get("verticalAlign")
        if vertical_align:
            cell.vertical_alignment = {
                "top": 0,
                "middle": 1,
                "bottom": 2
            }.get(vertical_align, 1)
    
    def _apply_header_style(self, cell, para):
        """为表头单元格应用默认样式"""
        # 设置灰色背景 (#F2F2F2)
        self._set_cell_background(cell, RGBColor(242, 242, 242))
        
        # 设置文本居中（如果没有其他对齐方式）
        if not para.alignment:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 设置文本加粗
        for run in para.runs:
            run.bold = True
    
    def _set_cell_background(self, cell, color: RGBColor):
        """设置单元格背景色"""
        try:
            shading_elm = OxmlElement('w:shd')
            # RGBColor 对象需要转换为十六进制字符串
            color_hex = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}" if isinstance(color, tuple) else f"{color.r:02X}{color.g:02X}{color.b:02X}"
            shading_elm.set(qn('w:fill'), color_hex)
            cell._element.get_or_add_tcPr().append(shading_elm)
        except Exception as e:
            logger.error(f"设置单元格背景色失败: {e}")
    
    def _set_paragraph_shading(self, para, color: RGBColor):
        """设置段落背景色"""
        try:
            pPr = para._element.get_or_add_pPr()
            shading = OxmlElement('w:shd')
            # RGBColor 对象需要转换为十六进制字符串
            color_hex = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}" if isinstance(color, tuple) else f"{color.r:02X}{color.g:02X}{color.b:02X}"
            shading.set(qn('w:fill'), color_hex)
            pPr.append(shading)
        except Exception as e:
            logger.error(f"设置段落背景色失败: {e}")
    
    def _set_run_background(self, run, color: RGBColor):
        """设置 run 背景色"""
        try:
            rPr = run._element.get_or_add_rPr()
            shading = OxmlElement('w:shd')
            # RGBColor 对象需要转换为十六进制字符串
            color_hex = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}" if isinstance(color, tuple) else f"{color.r:02X}{color.g:02X}{color.b:02X}"
            shading.set(qn('w:fill'), color_hex)
            rPr.append(shading)
        except Exception as e:
            logger.error(f"设置 run 背景色失败: {e}")
    
    def _parse_color(self, color_str: str) -> Optional[Tuple[int, int, int]]:
        """
        解析颜色字符串为 RGB 元组
        
        Args:
            color_str: 颜色字符串(如 "#ff0000" 或 "rgb(255, 0, 0)")
            
        Returns:
            (r, g, b) 元组或 None
        """
        if not color_str:
            return None
        
        # 处理十六进制颜色
        hex_match = re.match(r'#([0-9a-fA-F]{6})', color_str)
        if hex_match:
            hex_color = hex_match.group(1)
            r = int(hex_color[0:2], 16)
            g = int(hex_color[2:4], 16)
            b = int(hex_color[4:6], 16)
            return (r, g, b)
        
        # 处理 rgb() 格式
        rgb_match = re.match(r'rgb\((\d+),\s*(\d+),\s*(\d+)\)', color_str)
        if rgb_match:
            r = int(rgb_match.group(1))
            g = int(rgb_match.group(2))
            b = int(rgb_match.group(3))
            return (r, g, b)
        
        # 处理 rgba() 格式(忽略 alpha)
        rgba_match = re.match(r'rgba\((\d+),\s*(\d+),\s*(\d+),\s*[\d.]+\)', color_str)
        if rgba_match:
            r = int(rgba_match.group(1))
            g = int(rgba_match.group(2))
            b = int(rgba_match.group(3))
            return (r, g, b)
        
        return None
    
    def _parse_font_size(self, size_str: Union[str, int, float]) -> Optional[float]:
        """
        解析字号字符串为 Pt 数值 (float)
        
        Args:
            size_str: 字号字符串(如 "16px", "10.5pt", "16")
            
        Returns:
            字号(Pt) 或 None
        """
        if not size_str:
            return None
        
        # 如果已经是数字，默认为 pt (因为现在的系统主要推崇 pt)
        if isinstance(size_str, (int, float)):
            return float(size_str)
        
        s = str(size_str).lower().strip()
        
        # 提取数值
        match = re.search(r'(\d+(\.\d+)?)', s)
        if not match:
            return None
            
        val = float(match.group(1))
        
        # 单位判断
        if 'px' in s:
            # px 转 pt (Web常用 1px = 0.75pt)
            return val * 0.75
        
        # 默认或明确为 pt -> 直接返回
        return val
