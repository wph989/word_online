"""
Word 文档导出器主类
"""

import io
import logging
from docx import Document
from docx.shared import RGBColor
from typing import Dict, Any, Optional

from .style_utils import (
    apply_paragraph_style,
    apply_cell_style,
    apply_page_margins,
    configure_heading_styles
)
from .block_processors import (
    add_paragraph,
    add_heading,
    add_table,
    add_image,
    add_code,
    add_divider
)
from .heading_numbering import create_heading_number_generator
from .auto_numbering import create_multilevel_numbering

logger = logging.getLogger(__name__)


class DocxExporter:
    """
    Word 文档导出器
    将结构化的 Content 和 StyleSheet 转换为 Word 文档
    """
    
    def __init__(
        self, 
        content: Dict[str, Any], 
        stylesheet: Dict[str, Any],
        document_settings: Optional[Dict[str, Any]] = None
    ):
        """
        初始化导出器
        
        Args:
            content: Content JSON
            stylesheet: StyleSheet JSON
            document_settings: 文档配置 (页边距、标题样式等)
        """
        self.content = content
        self.stylesheet = stylesheet
        self.document_settings = document_settings or {}
        self.doc = Document()
        
        # 应用页边距配置
        apply_page_margins(self.doc, self.document_settings)
        
        # 初始化标题编号（支持两种模式）
        numbering_config = self.document_settings.get('heading_numbering_style')
        
        # 检查是否使用自动编号
        if numbering_config and numbering_config.get('useAutoNumbering', False):
            # 模式1：Word自动编号
            self.auto_numbering_id = create_multilevel_numbering(self.doc, numbering_config)
            self.heading_number_generator = None
            if self.auto_numbering_id:
                style_name = numbering_config.get('style', 'style2')
                logger.info(f"启用Word自动编号: 样式={style_name}, numId={self.auto_numbering_id}")
        else:
            # 模式2：文本前缀编号
            self.auto_numbering_id = None
            self.heading_number_generator = create_heading_number_generator(numbering_config)
            if self.heading_number_generator:
                style_name = numbering_config.get('style', 'custom') if numbering_config else 'none'
                logger.info(f"启用文本前缀编号: 样式={style_name}")
        
        # 配置标题样式模板（如果使用自动编号，将编号应用到样式）
        configure_heading_styles(self.doc, self.document_settings, self.auto_numbering_id)
        
        self.style_map = self._build_style_map()
        self.cell_style_map = self._build_cell_style_map()
        self.column_width_map = self._build_column_width_map()
        
        logger.info(f"初始化 DocxExporter, Blocks 数量: {len(content.get('blocks', []))}")
        if document_settings:
            logger.info(f"应用文档配置: 页边距={document_settings.get('margin_top', 2.54)}cm(上)")
    
    def export(self) -> io.BytesIO:
        """
        执行导出,生成 Word 文档
        
        Returns:
            包含 .docx 文件的 BytesIO 对象
        """
        try:
            # 处理所有 Block
            for idx, block in enumerate(self.content.get("blocks", [])):
                try:
                    self._process_block(block)
                except Exception as e:
                    logger.error(f"处理 Block {idx} 失败: {e}", exc_info=True)
                    # 添加错误占位符
                    para = self.doc.add_paragraph(f"[导出错误: {block.get('type', 'unknown')}]")
                    para.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            
            # 保存到内存
            file_stream = io.BytesIO()
            self.doc.save(file_stream)
            file_stream.seek(0)
            
            logger.info(f"导出成功,文件大小: {file_stream.getbuffer().nbytes} 字节")
            return file_stream
            
        except Exception as e:
            logger.error(f"导出失败: {e}", exc_info=True)
            raise
    
    def _build_style_map(self) -> Dict[str, Dict[str, Any]]:
        """构建 Block ID 到样式的映射"""
        style_map = {}
        
        for rule in self.stylesheet.get("rules", []):
            target = rule.get("target", {})
            block_ids = target.get("blockIds", [])
            styles = rule.get("style", {})
            
            # 只处理 Block 级别的样式
            if target.get("blockType") in ["paragraph", "heading", "table"]:
                for block_id in block_ids:
                    if block_id not in style_map:
                        style_map[block_id] = {}
                    style_map[block_id].update(styles)
        
        return style_map
    
    def _build_cell_style_map(self) -> Dict[str, Dict[str, Any]]:
        """构建单元格 ID 到样式的映射"""
        cell_style_map = {}
        
        for rule in self.stylesheet.get("rules", []):
            target = rule.get("target", {})
            
            # 处理表格单元格样式 - 使用 blockIds
            if target.get("blockType") == "tableCell":
                block_ids = target.get("blockIds", [])
                styles = rule.get("style", {})
                
                for block_id in block_ids:
                    if block_id not in cell_style_map:
                        cell_style_map[block_id] = {}
                    cell_style_map[block_id].update(styles)
        
        # 调试日志
        if cell_style_map:
            logger.info(f"📊 导出统计: 单元格样式数量: {len(cell_style_map)}")
            for cell_id, styles in list(cell_style_map.items())[:5]:
                logger.info(f"  {cell_id}: {styles}")
        
        return cell_style_map
    
    def _build_column_width_map(self) -> Dict[str, Dict[int, str]]:
        """构建表格列宽映射
        
        Returns:
            Dict[table_id, Dict[column_index, width_value]]
        """
        column_width_map = {}
        
        for rule in self.stylesheet.get("rules", []):
            target = rule.get("target", {})
            
            # 处理表格列宽度样式
            if target.get("blockType") == "tableColumn":
                block_ids = target.get("blockIds", [])
                column_index = target.get("columnIndex")
                styles = rule.get("style", {})
                width = styles.get("width")
                
                if column_index is not None and width:
                    for block_id in block_ids:
                        if block_id not in column_width_map:
                            column_width_map[block_id] = {}
                        column_width_map[block_id][column_index] = width
        
        return column_width_map
    
    def _process_block(self, block: Dict[str, Any]):
        """
        处理单个 Block
        
        Args:
            block: Block 字典
        """
        block_type = block.get("type")
        
        if block_type == "paragraph":
            add_paragraph(self.doc, block, self.style_map)
        elif block_type == "heading":
            add_heading(self.doc, block, self.style_map, self.heading_number_generator, self.auto_numbering_id)
        elif block_type == "table":
            block_id = block.get("id", "")
            column_widths = self.column_width_map.get(block_id, {})
            add_table(self.doc, block, self.cell_style_map, column_widths)
        elif block_type == "image":
            add_image(self.doc, block)
        elif block_type == "code":
            add_code(self.doc, block)
        elif block_type == "divider":
            add_divider(self.doc, block)
        else:
            logger.warning(f"未知的 Block 类型: {block_type}")
