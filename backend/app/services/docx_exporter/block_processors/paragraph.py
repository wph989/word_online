"""
段落处理器
"""

import logging
from docx.shared import Inches
from typing import Dict, Any

from ..style_utils import apply_paragraph_style

logger = logging.getLogger(__name__)


def add_paragraph(doc, block: Dict[str, Any], style_map: Dict[str, Dict[str, Any]]):
    """添加段落"""
    text = block.get("text", "")
    marks = block.get("marks", [])
    block_id = block.get("id")
    attrs = block.get("attrs") or {}  # 处理 null 值
    
    # 创建段落
    para = doc.add_paragraph()
    
    # 应用列表样式
    list_type = attrs.get("listType")
    list_level = attrs.get("listLevel", 0)
    
    if list_type == "bullet":
        para.style = 'List Bullet'
        # 设置列表层级
        if list_level > 0:
            para.paragraph_format.left_indent = Inches(0.5 * (list_level + 1))
    elif list_type == "ordered":
        para.style = 'List Number'
        if list_level > 0:
            para.paragraph_format.left_indent = Inches(0.5 * (list_level + 1))
    
    # 添加文本和标记
    from ..parsers import apply_default_style_to_run, apply_marks_to_run
    add_text_with_marks(para, text, marks)
    
    # 应用段落样式
    apply_paragraph_style(para, block_id, style_map)


def add_text_with_marks(para, text: str, marks: list, default_style: Dict[str, Any] = None, filter_mark_types: list = None, inherit_paragraph_style: bool = False):
    """
    添加带标记的文本到段落
    
    Args:
        para: Word 段落对象
        text: 文本内容
        marks: 标记列表
        default_style: 默认样式 (如标题的字号、颜色等)
        filter_mark_types: 需要过滤的 mark 类型列表（这些 mark 将被忽略）
        inherit_paragraph_style: 是否继承段落样式(True 时不应用默认字体和字号)
    """
    from ..parsers import apply_default_style_to_run, apply_marks_to_run
    
    if not text:
        return
    
    # 过滤 marks
    if filter_mark_types:
        marks = [m for m in marks if m.get("type") not in filter_mark_types]
    
    
    if not marks:
        run = para.add_run(text)
        apply_default_style_to_run(run, default_style, inherit_paragraph_style)
        return
    
    # 构建字符级别的标记映射
    char_marks = [[] for _ in range(len(text))]
    for mark in marks:
        start, end = mark["range"]
        for i in range(start, min(end, len(text))):
            char_marks[i].append(mark)
    
    # 合并相同标记的连续字符
    segments = []
    current_start = 0
    current_marks = char_marks[0] if char_marks else []
    
    for i in range(1, len(text)):
        if char_marks[i] != current_marks:
            segments.append((current_start, i, current_marks))
            current_start = i
            current_marks = char_marks[i]
    
    # 添加最后一段
    segments.append((current_start, len(text), current_marks))
    
    # 为每个段落创建 run 并应用标记
    for start, end, segment_marks in segments:
        run = para.add_run(text[start:end])
        
        # 1. 先应用默认样式 (如果有传入 default_style，则使用它；否则使用全局默认)
        apply_default_style_to_run(run, default_style, inherit_paragraph_style)
        
        # 2. 再应用字符级标记（会覆盖默认值）
        apply_marks_to_run(run, segment_marks)
