"""
图片处理器
"""

import logging
import re
import base64
import requests
from io import BytesIO
from docx.shared import Inches, RGBColor
from typing import Dict, Any

logger = logging.getLogger(__name__)


def add_image(doc, block: Dict[str, Any]):
    """添加图片(支持本地路径、URL 和 base64)"""
    src = block.get("src", "")
    meta = block.get("meta", {})
    
    try:
        image_stream = None
        
        # 处理 base64 图片
        if src.startswith("data:image"):
            # 提取 base64 数据
            match = re.match(r'data:image/[^;]+;base64,(.+)', src)
            if match:
                base64_data = match.group(1)
                image_data = base64.b64decode(base64_data)
                image_stream = BytesIO(image_data)
        
        # 处理 URL 图片
        elif src.startswith("http://") or src.startswith("https://"):
            response = requests.get(src, timeout=10)
            response.raise_for_status()
            image_stream = BytesIO(response.content)
        
        # 处理本地路径
        else:
            image_stream = src
        
        # 添加图片
        if image_stream:
            width = meta.get("width")
            if width and isinstance(width, (int, float)):
                width_inches = Inches(width / 96)  # 假设 96 DPI
                doc.add_picture(image_stream, width=width_inches)
            else:
                doc.add_picture(image_stream)
            
            logger.info(f"成功添加图片: {src[:50]}...")
        else:
            raise ValueError("无法识别的图片源")
            
    except Exception as e:
        logger.error(f"添加图片失败: {e}")
        # 添加占位文本
        para = doc.add_paragraph(f"[图片加载失败: {src[:50]}...]")
        para.runs[0].italic = True
        para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
