"""
表格处理器
"""

import logging
from typing import Dict, Any

from docx.shared import Twips

from ..style_utils import apply_cell_style, apply_header_style

logger = logging.getLogger(__name__)

# 像素到 Twips 的转换比例 (1 px ≈ 15 twips，基于 96 DPI)
PX_TO_TWIPS = 15


def add_table(
    doc, 
    block: Dict[str, Any], 
    cell_style_map: Dict[str, Dict[str, Any]],
    column_widths: Dict[int, str] = None
):
    """添加表格(支持合并单元格、单元格样式和列宽)
    
    Args:
        doc: Word 文档对象
        block: 表格 Block 数据
        cell_style_map: 单元格样式映射
        column_widths: 列宽映射 {column_index: width_value}，宽度值为像素
    """
    table_data = block.get("data", {})
    rows = table_data.get("rows", 0)
    cols = table_data.get("cols", 0)
    cells = table_data.get("cells", [])
    merge_regions = table_data.get("mergeRegions", [])
    
    if rows == 0 or cols == 0:
        logger.warning("表格行数或列数为 0,跳过")
        return
    
    # 创建表格
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    # 应用列宽
    if column_widths:
        _apply_column_widths(table, column_widths, cols)
    
    # 构建单元格映射
    cell_map = {tuple(cell["cell"]): cell for cell in cells}
    
    # 填充单元格内容和样式
    for row_idx in range(rows):
        for col_idx in range(cols):
            cell_data = cell_map.get((row_idx, col_idx))
            if cell_data:
                cell = table.cell(row_idx, col_idx)
                content = cell_data.get("content", {})
                text = content.get("text", "")
                marks = content.get("marks", [])
                
                # 清空默认段落
                cell.text = ""
                para = cell.paragraphs[0]
                
                # 添加文本和标记
                from .paragraph import add_text_with_marks
                add_text_with_marks(para, text, marks)
                
                # 应用单元格样式 - 使用 cell-X-Y 格式的 ID
                cell_id = f"cell-{row_idx}-{col_idx}"
                apply_cell_style(cell, para, cell_id, cell_style_map)
    
    # 处理合并单元格
    for region in merge_regions:
        try:
            start_row, start_col = region["start"]
            end_row, end_col = region["end"]
            
            # 合并单元格
            start_cell = table.cell(start_row, start_col)
            end_cell = table.cell(end_row, end_col)
            start_cell.merge(end_cell)
        except Exception as e:
            logger.error(f"合并单元格失败: {e}")


def _apply_column_widths(table, column_widths: Dict[int, str], cols: int):
    """应用列宽到表格
    
    Args:
        table: Word 表格对象
        column_widths: 列宽映射 {column_index: width_value}，宽度值可以是数字或带单位的字符串（如 "219.5pt"）
        cols: 列数
    """
    import re
    
    def parse_width(width_str):
        """解析宽度值，支持带单位的字符串
        
        Args:
            width_str: 宽度字符串，如 "219.5pt", "100px", "150"
            
        Returns:
            float: 宽度的点数值（pt）
        """
        if isinstance(width_str, (int, float)):
            return float(width_str)
        
        # 提取数字部分
        match = re.match(r'([\d.]+)\s*(pt|px|%)?', str(width_str).strip())
        if match:
            value = float(match.group(1))
            unit = match.group(2) or 'px'  # 默认单位为 px
            
            # 转换为 pt
            if unit == 'pt':
                return value
            elif unit == 'px':
                # 1 px ≈ 0.75 pt (基于 96 DPI)
                return value * 0.75
            elif unit == '%':
                # 百分比转换为固定宽度（假设页面宽度为 595pt，A4纸宽度）
                return 595 * value / 100
        
        return None
    
    try:
        for col_idx in range(cols):
            width_str = column_widths.get(col_idx)
            if width_str:
                try:
                    # 解析宽度值
                    width_pt = parse_width(width_str)
                    
                    if width_pt is not None:
                        # pt 转换为 Twips (1 pt = 20 twips)
                        width_twips = int(width_pt * 20)
                        
                        # 设置每一行的单元格宽度（确保列宽一致）
                        for row in table.rows:
                            if col_idx < len(row.cells):
                                row.cells[col_idx].width = Twips(width_twips)
                        
                        logger.debug(f"应用列宽: 列{col_idx} = {width_str} -> {width_pt}pt -> {width_twips} twips")
                    else:
                        logger.warning(f"无法解析列宽 (列{col_idx}): {width_str}")
                except Exception as e:
                    logger.warning(f"解析列宽失败 (列{col_idx}): {width_str}, 错误: {e}")
    except Exception as e:
        logger.error(f"应用列宽失败: {e}")

