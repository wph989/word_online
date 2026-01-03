"""
综合测试:验证标题样式模板功能
测试场景:
1. 标题样式模板正确配置
2. 标题文本继承样式模板
3. 局部 mark 格式正确叠加
4. 修改样式模板后所有标题同步更新
"""

import sys
import os
sys.path.insert(0, os.path.abspath('..'))

from app.services.docx_exporter import DocxExporter
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import logging

logging.basicConfig(level=logging.INFO)

print("=" * 80)
print("综合测试:标题样式模板功能")
print("=" * 80)

# 测试数据
content = {
    "blocks": [
        {
            "id": "h1-1",
            "type": "heading",
            "level": 1,
            "text": "第一章 项目概述",
            "marks": []
        },
        {
            "id": "p1",
            "type": "paragraph",
            "text": "这是一个普通段落,使用默认字体和字号。",
            "marks": [],
            "attrs": {"listType": "none"}
        },
        {
            "id": "h2-1",
            "type": "heading",
            "level": 2,
            "text": "1.1 项目背景",
            "marks": []
        },
        {
            "id": "p2",
            "type": "paragraph",
            "text": "这是另一个普通段落。",
            "marks": [],
            "attrs": {"listType": "none"}
        },
        {
            "id": "h2-2",
            "type": "heading",
            "level": 2,
            "text": "1.2 重要提示",
            "marks": [
                {"type": "color", "value": "#ff0000", "range": [4, 8]}  # "重要提示" 为红色
            ]
        },
        {
            "id": "h1-2",
            "type": "heading",
            "level": 1,
            "text": "第二章 技术架构",
            "marks": [
                {"type": "bold", "range": [4, 8]},  # "技术架构" 额外加粗
                {"type": "color", "value": "#0066cc", "range": [4, 8]}  # "技术架构" 蓝色
            ]
        }
    ]
}

stylesheet = {
    "styleId": "test-style",
    "appliesTo": "chapter",
    "rules": []
}

# 文档配置
document_settings = {
    "heading_styles": {
        "h1": {
            "fontSize": 24,
            "fontFamily": "微软雅黑",
            "color": "#000000",
            "fontWeight": "bold",
            "marginTop": 16,
            "marginBottom": 12
        },
        "h2": {
            "fontSize": 18,
            "fontFamily": "微软雅黑",
            "color": "#333333",
            "fontWeight": "bold",
            "marginTop": 12,
            "marginBottom": 8
        }
    }
}

print("\n【测试场景 1】导出文档并验证样式模板")
print("-" * 80)

# 创建导出器
exporter = DocxExporter(content, stylesheet, document_settings)
file_stream = exporter.export()

# 保存文档
output_path = "test_comprehensive_heading.docx"
with open(output_path, "wb") as f:
    f.write(file_stream.read())

print(f"✅ 文档已导出: {output_path}")

# 验证文档
doc = Document(output_path)

print("\n【测试场景 2】验证标题样式模板配置")
print("-" * 80)

test_passed = True

# 验证 Heading 1
h1_style = doc.styles['Heading 1']
print("\nHeading 1 样式模板:")
print(f"  字体: {h1_style.font.name}")
print(f"  字号: {h1_style.font.size.pt if h1_style.font.size else 'None'} pt")
print(f"  颜色: #{h1_style.font.color.rgb if h1_style.font.color.rgb else 'None'}")
print(f"  加粗: {h1_style.font.bold}")

if h1_style.font.size.pt != 24:
    print(f"  ❌ 错误: 字号应为 24pt, 实际为 {h1_style.font.size.pt}pt")
    test_passed = False
else:
    print(f"  ✅ 字号正确")

# 验证 Heading 2
h2_style = doc.styles['Heading 2']
print("\nHeading 2 样式模板:")
print(f"  字体: {h2_style.font.name}")
print(f"  字号: {h2_style.font.size.pt if h2_style.font.size else 'None'} pt")
print(f"  颜色: #{h2_style.font.color.rgb if h2_style.font.color.rgb else 'None'}")
print(f"  加粗: {h2_style.font.bold}")

if h2_style.font.size.pt != 18:
    print(f"  ❌ 错误: 字号应为 18pt, 实际为 {h2_style.font.size.pt}pt")
    test_passed = False
else:
    print(f"  ✅ 字号正确")

print("\n【测试场景 3】验证标题文本继承样式模板")
print("-" * 80)

# 查找标题段落
heading_paras = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]

for para in heading_paras:
    print(f"\n{para.style.name}: {para.text}")
    
    # 检查 run 是否没有设置字体和字号(应该继承样式模板)
    for i, run in enumerate(para.runs):
        has_font = run.font.name is not None
        has_size = run.font.size is not None
        
        # 对于没有 mark 的 run,不应该有字体和字号设置
        if not run.font.color.rgb and not run.bold:
            if has_font or has_size:
                print(f"  ❌ Run {i+1} 不应该设置字体或字号(应继承样式模板)")
                test_passed = False
            else:
                print(f"  ✅ Run {i+1} 正确继承样式模板")

print("\n【测试场景 4】验证局部 mark 格式叠加")
print("-" * 80)

# 查找带 mark 的标题
para_with_red = [p for p in doc.paragraphs if "重要提示" in p.text][0]
print(f"\n段落: {para_with_red.text}")
for i, run in enumerate(para_with_red.runs):
    print(f"  Run {i+1}: '{run.text}'")
    if run.font.color.rgb:
        print(f"    颜色: {run.font.color.rgb}")
        if run.font.color.rgb == RGBColor(255, 0, 0):
            print(f"    ✅ 红色正确")
        else:
            print(f"    ❌ 颜色错误")
            test_passed = False

para_with_blue = [p for p in doc.paragraphs if "技术架构" in p.text][0]
print(f"\n段落: {para_with_blue.text}")
for i, run in enumerate(para_with_blue.runs):
    print(f"  Run {i+1}: '{run.text}'")
    if run.font.color.rgb:
        print(f"    颜色: {run.font.color.rgb}")
        if run.font.color.rgb == RGBColor(0, 102, 204):
            print(f"    ✅ 蓝色正确")
        else:
            print(f"    ❌ 颜色错误")
            test_passed = False
    if run.bold:
        print(f"    ✅ 加粗正确")

print("\n【测试场景 5】验证普通段落使用默认样式")
print("-" * 80)

normal_paras = [p for p in doc.paragraphs if p.style.name == 'Normal']
for para in normal_paras[:2]:  # 只检查前两个
    print(f"\n段落: {para.text}")
    for i, run in enumerate(para.runs):
        print(f"  Run {i+1}:")
        print(f"    字体: {run.font.name}")
        print(f"    字号: {run.font.size.pt if run.font.size else 'None'} pt")
        
        # 普通段落应该有默认字体和字号
        if run.font.name and run.font.size:
            print(f"    ✅ 正确应用默认样式")
        else:
            print(f"    ❌ 缺少默认样式")
            test_passed = False

print("\n" + "=" * 80)
if test_passed:
    print("🎉 所有测试通过!")
else:
    print("❌ 部分测试失败")
print("=" * 80)

print("\n【手动验证步骤】")
print("请在 Word 中打开 test_comprehensive_heading.docx 并执行以下操作:")
print("1. 右键点击任意一级标题 → 修改样式 → 修改字号为 28pt")
print("2. 观察所有一级标题是否同步更新为 28pt")
print("3. 验证带红色/蓝色的文字是否保持颜色不变")
print("4. 这证明了标题使用样式模板,而 mark 格式正确叠加")
